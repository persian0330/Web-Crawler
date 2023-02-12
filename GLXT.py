'''
This code is used for simplify the precess of submitting form data on an certain website.
Implemented:
    Simulate the login process
    Send GET request to get web page information after login
    Send POST request to the web page to submit the prepared form data
'''


# -*- coding: utf-8 -*
import requests
import json
from win32com import client as wc
from docx import Document

#增加报告 simulate a post request: Add a report to the website
def add(session,headers,renwuId,bjdId,baseId):

    add_data = {
        'modelId': '21001',
        'shebeizhongleidaima': '2000',
        'shebeileibiedaima': '2100',
        'shebeipinzhongdaima': '2110',
        'renwuId': renwuId,  # 查询可得
        'bjdId': bjdId,  # 查询可得
        'baseId': baseId,  # 查询可得
        'ysjlId': '',
        'quhuamingcheng': '开发区',  #查询可得
        'quhuadaima': '370698',
        'jianyanxiangmu': '2100',
        'version': '2.0',  #版本号固定
        'jigouhezhunzhenghao': 'TS7110434-2024',#?
        'bgbh1': dict['jilubianhao'][0],
        'bgbh2': dict['jilubianhao'][1],
        'bgbh3': dict['jilubianhao'][2],
        'bgbh4': dict['jilubianhao'][3],
        'baogaobianhao': dict['jilubianhao'][0]+'-'+dict['jilubianhao'][1]+'-'+dict['jilubianhao'][2]+'-'+dict['jilubianhao'][3],
        'zhucedaima': dict['zhucedaima'],
        'shebeidaima': dict['shebeidaima'],
        'shebeimingcheng': dict['shebeimingcheng'],
        'jianyanleibie': dict['jianyanleibie'],
        'rongqileibie': dict['rongqileibie'],
        'danweineibubianhao': dict['danweineibianhao'],
        'shiyongdengjibianhao': dict['shiyongdengjizhengbianhao'],
        'zhizaodanwei': dict['zhizaodanwei'],
        'shigongdanwei': dict['anzhuangdanwei'],
        'shiyongdanwei': dict['shiyongdanwei'],
        'shiyongdanweiDizhi': dict['shiyongdanweidizhi'],
        'anzhuangdidian': dict['shebeishiyongdidian'],
        'shiyongdanweiCode': dict['shiyongdanweizuzhijigoudaima'],
        'shiyongdanweiYoubian': dict['youzhengbianma'],
        'anquanguanlirenyuan': dict['anquanguanlirenyuan'],
        'anquanguanlidianhua': dict['lianxidianhua'],
        'shejishiyongnianxian': dict['shejishiyongnianxian'],
        'touyongriqi': dict['tourishiyongriqi'],
        'zhutijiegouxingshi': dict['zhutijiegouxingshi'],
        'jiegou': '夹套',
        'yunxingzhuangtai': dict['yunxingzhuangtai'],
        'rongji': dict['rongji'],
        'neijing': dict['rongqineijing'],
        'shejiyali': dict['shejiyaliketi']+'；'+dict['shejiyalijiatao'],
        'shejiwendu': dict['shejiwenduketi']+'；'+dict['shejiwendujiatao'],
        'shiyongyali': dict['gongzuoyaliketi']+'；'+dict['gongzuoyalijiatao'],
        'shiyongwendu': dict['gongzuowenduketi']+'；'+dict['gongzuowendujiatao'],
        'gongzuojiezhi': dict['jiezhiketi']+'；'+dict['jiezhijiatao'],
        'yali': dict['shiyongcanshuyali'],
        'wendu': dict['shiyongcanshuwendu'],
        'jiezhi': dict['shiyongcanshujiezhi'],
        'qita': dict['shiyongcanshuqita'],
        'jianyanyiju': dict['jianyanyiju'],
        'wentijiqichuli': dict['wentijiqichuli'],#?
        'jianyanjielun': '符合要求',  #?
        'anquanzhuangkuangdengji': '2级',  #?
        'jianyanriqi': '2020-12-01',#?
        'jianyanjieshuriqi': '2020-12-01',#?
        'xiacijianyanriqi': '2025-12-12',#?
        'bianzhiriqi': '2020-12-30',#?
        'beizhu': dict['shuoming'],
        'bianzhiren': '李培尧',#?
        'shejidanwei': dict['shejidanwei'],
        'shejiriqi': dict['shejiriqi'],
        'chanpinbiaozhun': dict['chanpinbiaozhun'],
        'rongqituhao': dict['rongqituhao'],
        'zzdw2': dict['zhizaodanwei'],
        'zhizaoriqi': dict['zhizaoriqi'],
        'chanpinbianhao': dict['chanpinbianhao'],
        'shangcijianyanriqi': dict['shangcijianyanriqi'],
        'sydjz2': dict['shiyongdengjizhengbianhao'],
        'anzhuangxingshi': dict['anzhuangxingshi'],
        'zhizuoxingshi': dict['zhizuoxingshi'],
        'baowenjuerefangshi': dict['baowenjuerefangshi'],
        'rongqihuanremianji': dict['rongji'],
        'rongqineijing': dict['rongqineijing'],
        'gao': dict['gaochang'],
        'chongzhuangxishu': dict['zuidayunxuchongzhuangliang'],
        'waiketongticailiao': '壳体',#?
        'ketishejiyali': dict['shejiyaliketi'],
        'ketigongzuoyali': dict['gongzuoyaliketi'],
        'waiketongtihoudu': '夹套',#?
        'guanchengshejiyali': dict['shejiyalijiatao'],
        'jiataogongzuoyali': dict['gongzuoyalijiatao'],
        'ketishejiwendu': dict['shejiwenduketi'],
        'ketigongzuowendu': dict['gongzuowenduketi'],
        'jiataoshejiwendu': dict['shejiwendujiatao'],
        'jiataogongzuowendu': dict['gongzuowendujiatao'],
        'tongtifushiyudu': dict['fushiyudutongti'],
        'ketijiezhi': dict['jiezhiketi'],
        'fengtoufushiyudu': dict['fushiyudufengtou'],
        'jiataojiezhi': dict['jiezhijiatao'],
        'tongticaizhi': dict['caizhitongti'],
        'tongtihoudu': dict['houdutongti'],
        'fengtoucaizhi': dict['caizhifengtou'],
        'fengtouhoudu': dict['houdufengtou'],
        'waikefengtoucailiao': '夹套',#?
        'jiataocaizhi': dict['caizhijiatao'],
        'jiataohoudu': dict['jiatao'],
        'chenlicaizhi': dict['caizhichenli'],
        'chenlihoudu': dict['chenli'],
        'ziliaoshenchaqingkuang': dict['ziliaoshenchaqingquang'],
        'shangcianquandengji': '2级',#?
        'shagnciwenti': '无',#?
        'hongguanjianyanJieguo': '合格',
        'jyxm': str(jyxm)
    }

    url_add = 'http://yantaisei.sdsei.org.cn:8092/YTSE/2000/add?modeId=21001'

    resp_add = session.post(url_add, headers=headers, data=add_data)

    temp = json.loads(resp_add.text)

    print('报告已添加,报告ID：'+temp['report']['annalid'])

    return temp['report']['annalid']

#增加壁厚测定报告 simulate a post request: Add a certain report to the website
def add_bhcd(session,headers,annalid):

    #获得增加壁厚测定的reportID
    url_get_jyfaID = 'http://yantaisei.sdsei.org.cn:8092/YTSE/2000/jyfaList?reportId='+annalid

    data_reportID = {'reportId': annalid}

    resp = session.post(url_get_jyfaID,headers=headers,data=data_reportID)

    temp = json.loads(resp.text)
    #temp['rows'][0]#无损检测
    #temp['rows'][1]为分项检测,temp['rows'][1]['children']为分项检测下所有子检测,temp['rows'][1]['children'][0])为壁厚测定

    reportID = temp['rows'][1]['children'][0]['id']#获得增加壁厚测定的报告ID
    print(reportID)

    #增加壁厚测定
    url_add_jyfa = 'http://yantaisei.sdsei.org.cn:8092/YTSE/2000/addJyfa'

    data_jyfa = {
        'reportId': reportID,
        'type': '壁厚测定',
        'reportid': annalid
    }

    resp = session.post(url_add_jyfa,headers=headers,data=data_jyfa)

    #修改壁厚测定
    #获取新添加的壁厚测定报告ID

    url_get_bhcdID = 'http://yantaisei.sdsei.org.cn:8092/YTSE/2000/viewqita?reportId='+annalid
    resp = session.get(url_get_bhcdID,headers=headers)
    temp = json.loads(resp.text)

    #获得壁厚测定报告ID
    annalid = temp['rows'][0]['annalid']
    print('壁厚测定报告ID为'+annalid)

    url_edit_bhcd = 'http://yantaisei.sdsei.org.cn:8092/YTSE/fx/saveBhcd'
    
    data_bhcd = {
        'annalid': annalid,
        'baogaobianhao': dict['jilubianhao'][0]+'-'+dict['jilubianhao'][1]+'-'+dict['jilubianhao'][2]+'-'+dict['jilubianhao'][3],
        'celiangyiqixinghao': dict['celiangyiqixinghao'],
        'celiangyiqibianhao': dict['celiangyiqibianhao'],
        'celiangyiqijingdu': dict['celiangyiqijingdu'],
        'ouheji': dict['yuheji'],
        'tongtimingyihoudu': dict['mingyihoudutongti'],
        'tongtizuixiaobihou': dict['shicehoudutongti'],
        'fengtoumingyihoudu': dict['mingyihoudufengtou'],
        'fengtouzuixiaobihou': dict['shicehoudufengtou'],
        'costomA': dict['mingyihoudubuchong1_name'],
        'costomAValue': dict['mingyihoudubuchong1'],
        'costomB': dict['shicehoudubuchong1_name'],
        'costomBValue': dict['shicehoudubuchong1'],
        'costomC': dict['mingyihoudubuchong2_name'],
        'costomCValue': dict['mingyihoudubuchong2'],
        'costomD': dict['shicehoudubuchong2_name'],
        'costomDValue': dict['shicehoudubuchong2'],
        'biaomianzhuangkuang': dict['biaomianqingkuang'],
        'shicedianshu': dict['shicedianshu'],
        'jianyanjielun': dict['bhcdjianyanjielun'],
        'jianyanriqi': '2020-12-01',
        'bhcdfb': '[{"bianhaoone":1,"bianhaotwo":2,"bianhaothree":3,"bianhaofour":4,"bianhaofive":5,"bianhaosix":6,"houduone":"13.8","houdutwo":"13.9","houduthree":"13.9","houdufour":"13.8","houdufive":"13.8","houdusix":"13.9"},{"bianhaoone":7,"bianhaotwo":8,"bianhaothree":9,"bianhaofour":10,"bianhaofive":11,"bianhaosix":12,"houduone":"13.9","houdutwo":"13.9","houduthree":"13.8","houdufour":"13.9","houdufive":"13.9","houdusix":"13.9"},{"bianhaoone":13,"bianhaotwo":14,"bianhaothree":15,"bianhaofour":16,"bianhaofive":17,"bianhaosix":18,"houduone":"","houdutwo":"","houduthree":"","houdufour":"","houdufive":"","houdusix":""},{"bianhaoone":19,"bianhaotwo":20,"bianhaothree":21,"bianhaofour":22,"bianhaofive":23,"bianhaosix":24,"houduone":"","houdutwo":"","houduthree":"","houdufour":"","houdufive":"","houdusix":""}]'
    }

    resp_add_bhcd = session.post(url_edit_bhcd, headers=headers, data=data_bhcd)

    print('壁厚测定已添加')

#转换格式convert type
def doc2docx(path):
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(path)  # 目标路径下的文件
    doc.SaveAs(path+'x', 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()
    return path+'x'

#加载并读取doc文档 load and read a doc
def load(path):

    document = Document(path)
    tables = document.tables #获取文件中的表格集

    dict = {}

    # 获取记录编号
    bianhao_start = document.paragraphs[1].text.find('记录编号')
    dict['jilubianhao']=document.paragraphs[1].text[bianhao_start+5:].split('-',3)

    #表1
    dict['shebeimingcheng']=tables[0].cell(0,2).text
    dict['jianyanleibie']=tables[0].cell(0,8).text
    dict['rongqileibie']=tables[0].cell(1,2).text
    dict['shebeidaima']=tables[0].cell(3,2).text
    dict['danweineibianhao']=tables[0].cell(2,2).text
    dict['shiyongdengjizhengbianhao']=tables[0].cell(2,8).text
    dict['zhucedaima']=tables[0].cell(3,2).text
    dict['shiyongdanwei']=tables[0].cell(4,2).text
    dict['shiyongdanweidizhi']=tables[0].cell(5,2).text
    dict['shebeishiyongdidian']=tables[0].cell(6,2).text
    dict['yunxingzhuangtai']=tables[0].cell(6,8).text
    dict['shiyongdanweizuzhijigoudaima']=tables[0].cell(7,2).text
    dict['youzhengbianma']=tables[0].cell(7,8).text
    dict['anquanguanlirenyuan']=tables[0].cell(8,2).text
    dict['lianxidianhua']=tables[0].cell(8,8).text
    dict['jianyanyiju']=tables[0].cell(9,2).text
    dict['wentijiqichuli']=tables[0].cell(10,2).text
    dict['yalirongqianquanzhuangkuang']=tables[0].cell(11,2).text
    dict['fuheyaoqiu']=tables[0].cell(12,2).text
    dict['shiyongcanshuyali']=tables[0].cell(13,4).text
    dict['shiyongcanshuwendu']=tables[0].cell(13,7).text
    dict['shiyongcanshujiezhi']=tables[0].cell(14,4).text
    dict['shiyongcanshuqita']=tables[0].cell(14,7).text
    dict['xiacidingqijianyanriqi']=tables[0].cell(15,2).text
    dict['shuoming']=tables[0].cell(16,1).text
    #表2
    dict['shejidanwei']=tables[1].cell(0,3).text
    dict['shejiriqi']=tables[1].cell(1,3).text
    dict['chanpinbiaozhun']=tables[1].cell(1,8).text
    dict['rongqituhao']=tables[1].cell(2,3).text
    dict['shejishiyongnianxian']=tables[1].cell(2,8).text
    dict['zhizaodanwei']=tables[1].cell(3,3).text
    dict['zhizaoriqi']=tables[1].cell(4,3).text
    dict['chanpinbianhao']=tables[1].cell(4,8).text
    dict['anzhuangdanwei']=tables[1].cell(5,3).text
    dict['tourishiyongriqi']=tables[1].cell(6,3).text
    dict['shangcijianyanriqi']=tables[1].cell(6,8).text
    dict['shebeidanma']=tables[1].cell(7,3).text
    dict['zhutijiegouxingshi']=tables[1].cell(8,3).text
    dict['anzhuangxingshi']=tables[1].cell(8,8).text
    dict['zhizuoxingshi']=tables[1].cell(9,3).text
    dict['baowenjuerefangshi']=tables[1].cell(9,8).text
    dict['rongji']=tables[1].cell(10,3).text
    dict['rongqineijing']=tables[1].cell(10,8).text
    dict['gaochang']=tables[1].cell(11,3).text
    dict['zuidayunxuchongzhuangliang']=tables[1].cell(11,8).text
    dict['shejiyaliketi']=tables[1].cell(12,3).text
    dict['gongzuoyaliketi']=tables[1].cell(12,8).text
    dict['shejiyalijiatao']=tables[1].cell(13,3).text
    dict['gongzuoyalijiatao']=tables[1].cell(13,8).text
    dict['shejiwenduketi']=tables[1].cell(14,3).text
    dict['gongzuowenduketi']=tables[1].cell(14,8).text
    dict['shejiwendujiatao']=tables[1].cell(15,3).text
    dict['gongzuowendujiatao']=tables[1].cell(15,8).text
    dict['fushiyudutongti']=tables[1].cell(16,3).text
    dict['jiezhiketi']=tables[1].cell(16,8).text
    dict['fushiyudufengtou']=tables[1].cell(17,3).text
    dict['jiezhijiatao']=tables[1].cell(17,8).text
    dict['caizhitongti']=tables[1].cell(18,3).text
    dict['houdutongti']=tables[1].cell(18,8).text
    dict['caizhifengtou']=tables[1].cell(19,3).text
    dict['houdufengtou']=tables[1].cell(19,8).text
    dict['caizhijiatao']=tables[1].cell(20,3).text
    dict['jiatao']=tables[1].cell(20,8).text
    dict['caizhichenli']=tables[1].cell(21,3).text
    dict['chenli']=tables[1].cell(21,8).text
    dict['ziliaoshenchaqingquang']=tables[1].cell(22,2).text
    dict['shangcidingqijianyanwentijizai']=tables[1].cell(23,2).text
    #表3
    for i in range(27):
        dict[i+1]=tables[2].cell(i+1,4).text
        dict[i+1+100]=tables[2].cell(i+1,5).text
    #表4
    dict['celiangyiqixinghao']=tables[3].cell(0,4).text
    dict['celiangyiqibianhao'] = tables[3].cell(0,11).text
    dict['celiangyiqijingdu']=tables[3].cell(1,4).text
    dict['yuheji'] = tables[3].cell(1,11).text
    dict['mingyihoudutongti'] = tables[3].cell(2,4).text
    dict['shicehoudutongti'] = tables[3].cell(2,11).text
    dict['mingyihoudufengtou'] = tables[3].cell(3,4).text
    dict['shicehoudufengtou'] = tables[3].cell(3,11).text
    dict['mingyihoudubuchong1_name'] = tables[3].cell(4,2).text
    dict['mingyihoudubuchong1'] = tables[3].cell(4,4).text
    dict['shicehoudubuchong1_name'] = tables[3].cell(4,9).text
    dict['shicehoudubuchong1'] = tables[3].cell(4,11).text
    dict['mingyihoudubuchong2_name'] = tables[3].cell(5,2).text
    dict['mingyihoudubuchong2'] = tables[3].cell(5,4).text
    dict['shicehoudubuchong2_name'] = tables[3].cell(5,9).text
    dict['shicehoudubuchong2'] = tables[3].cell(5,11).text
    dict['biaomianqingkuang'] = tables[3].cell(6,4).text
    dict['shicedianshu'] = tables[3].cell(6,11).text
    for i in range(8):
        dict[201+i*6] = tables[3].cell(10+i,1).text
        dict[202+i*6] = tables[3].cell(10+i,4).text
        dict[203+i*6] = tables[3].cell(10+i,6).text
        dict[204+i*6] = tables[3].cell(10+i,8).text
        dict[205+i*6] = tables[3].cell(10+i,11).text
        dict[206+i*6] = tables[3].cell(10+i,13).text
    dict['bhcdjianyanjielun'] = tables[3].cell(18,0).text
    #表5
    #print(dict)
    print('已读取doc数据')
    return dict

#登录账号 url to login
url_login = 'http://yantaisei.sdsei.org.cn:8092/YTSE/user/login'
url_rw = 'http://yantaisei.sdsei.org.cn:8092/YTSE/renwu/rwlist'
#数据头文件 header
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.79 Safari/537.36',
}
#账号登录 login data
login_data = {
    'username': 'lipy',
    'password': '123'
}

session = requests.session()
session.post(url_login,headers=headers,data=login_data)

#读取报告 Get information from a doc

dict = load(doc2docx('C:/Users/Li/Desktop/编程/报告/1597——F00918.doc'))

jyxm = [
    {"annalid": 12139, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "6,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "封头型式",
         "jianyanjieguo": dict[1], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 1,
         "pdfbianhao": "6,1,0,0", "jianyanneirong": "封头型式", "xiangxineirong": "封头型式", "beizhu": dict[101]},
    {"annalid": 12140, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "封头与筒体的连接",
         "jianyanjieguo": dict[2], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 2,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "封头与筒体的连接", "xiangxineirong": "封头与筒体的连接", "beizhu": dict[102]},
    {"annalid": 12141, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "开孔位置及补强",
         "jianyanjieguo": dict[3], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 3,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "开孔位置及补强", "xiangxineirong": "开孔位置及补强", "beizhu": dict[103]},
    {"annalid": 12142, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "纵/（环）焊缝的布置及型式",
         "jianyanjieguo": dict[4], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 4,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "纵/（环）焊缝的布置及型式", "xiangxineirong": "纵/（环）焊缝的布置及型式",
         "beizhu": dict[104]},
    {"annalid": 12143, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "支承或者支座的型式与布置",
         "jianyanjieguo": dict[5], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 5,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "支承或者支座的型式与布置", "xiangxineirong": "支承或者支座的型式与布置",
         "beizhu": dict[105]},
    {"annalid": 12144, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "结构检验", "jianyanmu": "排放（疏水、排污）装置的设置",
         "jianyanjieguo": dict[6], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 6,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "排放（疏水、排污）装置的设置", "xiangxineirong": "排放（疏水、排污）装置的设置",
         "beizhu": dict[106]},
    {"annalid": 12145, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "5,1,0,0", "xiangmuleibie": "", "jianyanxiang": "几何尺寸检验", "jianyanmu": "筒体同一断面上最大内径与最小内径之差",
         "jianyanjieguo": dict[7], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 7,
         "pdfbianhao": "5,1,0,0", "jianyanneirong": "筒体同一断面上最大内径与最小内径之差", "xiangxineirong": "筒体同一断面上最大内径与最小内径之差",
         "beizhu": dict[107]},
    {"annalid": 12146, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "几何尺寸检验", "jianyanmu": "纵/（环）焊缝最大对口错边量",
         "jianyanjieguo": dict[8], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 8,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "纵/（环）焊缝最大对口错边量", "xiangxineirong": "纵/（环）焊缝最大对口错边量",
         "beizhu": dict[108]},
    {"annalid": 12147, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "几何尺寸检验", "jianyanmu": "纵/（环）焊缝最大棱角度",
         "jianyanjieguo": dict[9], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 9,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "纵/（环）焊缝最大棱角度", "xiangxineirong": "纵/（环）焊缝最大棱角度",
         "beizhu": dict[109]},
    {"annalid": 12148, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "几何尺寸检验", "jianyanmu": "纵/（环）焊缝最大咬边",
         "jianyanjieguo": dict[10], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 10,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "纵/（环）焊缝最大咬边", "xiangxineirong": "纵/（环）焊缝最大咬边",
         "beizhu": dict[110]},
    {"annalid": 12149, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "几何尺寸检验", "jianyanmu": "纵/（环）焊缝最大余高",
         "jianyanjieguo": dict[11], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 11,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "纵/（环）焊缝最大余高", "xiangxineirong": "纵/（环）焊缝最大余高",
         "beizhu": dict[111]},
    {"annalid": 12150, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "10,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "铭牌和标志",
         "jianyanjieguo": dict[12], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 12,
         "pdfbianhao": "10,1,0,0", "jianyanneirong": "铭牌和标志", "xiangxineirong": "铭牌和标志", "beizhu": dict[112]},
    {"annalid": 12151, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "内外表面腐蚀",
         "jianyanjieguo": dict[13], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 13,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "内外表面腐蚀", "xiangxineirong": "内外表面腐蚀", "beizhu": dict[113]},
    {"annalid": 12152, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验",
         "jianyanmu": "裂纹、泄漏、鼓包、变形、机械接触损伤、过热", "jianyanjieguo": dict[14], "jianyanjielun": "合格", "gongzuojianzheng": "",
         "querenfangshi": "", "sort": 14, "pdfbianhao": "0,1,0,0", "jianyanneirong": "裂纹、泄漏、鼓包、变形、机械接触损伤、过热",
         "xiangxineirong": "裂纹、泄漏、鼓包、变形、机械接触损伤、过热", "beizhu": dict[114]},
    {"annalid": 12153, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "工卡具焊迹、电弧灼伤",
         "jianyanjieguo": dict[15], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 15,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "工卡具焊迹、电弧灼伤", "xiangxineirong": "工卡具焊迹、电弧灼伤", "beizhu": dict[115]},
    {"annalid": 12154, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "法兰、密封面及其紧固螺栓",
         "jianyanjieguo": dict[16], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 16,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "法兰、密封面及其紧固螺栓", "xiangxineirong": "法兰、密封面及其紧固螺栓",
         "beizhu": dict[116]},
    {"annalid": 12155, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "支承、支座或者基础的下沉、倾斜、开裂",
         "jianyanjieguo": dict[17], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 17,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "支承、支座或者基础的下沉、倾斜、开裂", "xiangxineirong": "支承、支座或者基础的下沉、倾斜、开裂",
         "beizhu": dict[117]},
    {"annalid": 12156, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "地脚螺栓",
         "jianyanjieguo": dict[18], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 18,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "地脚螺栓", "xiangxineirong": "地脚螺栓", "beizhu": dict[118]},
    {"annalid": 12157, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "直立容器和球形容器支柱的铅垂度",
         "jianyanjieguo": dict[19], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 19,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "直立容器和球形容器支柱的铅垂度", "xiangxineirong": "直立容器和球形容器支柱的铅垂度",
         "beizhu": dict[119]},
    {"annalid": 12158, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验", "jianyanmu": "多支座卧式容器的支座膨胀孔",
         "jianyanjieguo": dict[20], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 20,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "多支座卧式容器的支座膨胀孔", "xiangxineirong": "多支座卧式容器的支座膨胀孔",
         "beizhu": dict[120]},
    {"annalid": 12159, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "壳体外观检验",
         "jianyanmu": "排放（疏水、排污）装置和泄漏信号指示孔的堵塞、腐蚀、沉积物", "jianyanjieguo": dict[21], "jianyanjielun": "合格",
         "gongzuojianzheng": "", "querenfangshi": "", "sort": 21, "pdfbianhao": "0,1,0,0",
         "jianyanneirong": "排放（疏水、排污）装置和泄漏信号指示孔的堵塞、腐蚀、沉积物", "xiangxineirong": "排放（疏水、排污）装置和泄漏信号指示孔的堵塞、腐蚀、沉积物",
         "beizhu": dict[121]},
    {"annalid": 12160, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "3,1,0,0", "xiangmuleibie": "", "jianyanxiang": "隔热层、衬里层检验",
         "jianyanmu": "隔热层破损、脱落、潮湿及层下腐蚀、裂纹", "jianyanjieguo": dict[22], "jianyanjielun": "合格", "gongzuojianzheng": "",
         "querenfangshi": "", "sort": 22, "pdfbianhao": "3,1,0,0", "jianyanneirong": "隔热层破损、脱落、潮湿及层下腐蚀、裂纹",
         "xiangxineirong": "隔热层破损、脱落、潮湿及层下腐蚀、裂纹", "beizhu": dict[122]},
    {"annalid": 12161, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "隔热层、衬里层检验",
         "jianyanmu": "衬里层的破损、腐蚀、裂纹、脱落及检查孔介质流出情况", "jianyanjieguo": dict[23], "jianyanjielun": "合格",
         "gongzuojianzheng": "", "querenfangshi": "", "sort": 23, "pdfbianhao": "0,1,0,0",
         "jianyanneirong": "衬里层的破损、腐蚀、裂纹、脱落及检查孔介质流出情况", "xiangxineirong": "衬里层的破损、腐蚀、裂纹、脱落及检查孔介质流出情况",
         "beizhu": dict[123]},
    {"annalid": 12162, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "隔热层、衬里层检验", "jianyanmu": "堆焊层的龟裂、剥离和脱落",
         "jianyanjieguo": dict[24], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 24,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "堆焊层的龟裂、剥离和脱落", "xiangxineirong": "堆焊层的龟裂、剥离和脱落",
         "beizhu": dict[124]},
    {"annalid": 12163, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "3,1,0,0", "xiangmuleibie": "", "jianyanxiang": "其他检验", "jianyanmu": "夹层真空度",
         "jianyanjieguo": dict[25], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 25,
         "pdfbianhao": "3,1,0,0", "jianyanneirong": "夹层真空度", "xiangxineirong": "夹层真空度", "beizhu": dict[125]},
    {"annalid": 12164, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "其他检验", "jianyanmu": "日蒸发率",
         "jianyanjieguo": dict[26], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 26,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "日蒸发率", "xiangxineirong": "日蒸发率", "beizhu": dict[126]},
    {"annalid": 12165, "modelId": "21001", "shebeizhongleidaima": "2000", "jianyanleibie": "DJ",
         "xiangmubianhao": "0,1,0,0", "xiangmuleibie": "", "jianyanxiang": "其他检验", "jianyanmu": "快开门式压力容器安全联锁功能",
         "jianyanjieguo": dict[27], "jianyanjielun": "合格", "gongzuojianzheng": "", "querenfangshi": "", "sort": 27,
         "pdfbianhao": "0,1,0,0", "jianyanneirong": "快开门式压力容器安全联锁功能", "xiangxineirong": "快开门式压力容器安全联锁功能",
         "beizhu": dict[127]}
    ]

print('使用登记证编号：'+dict['shiyongdengjizhengbianhao'][-5:])

#查找报告+增加报告

select_data = {
    'zhucedaima':'',
    'shebeizhongleidaima':'',
    'jianyanleibie':'',
    'shiyongdengjibianhao':dict['shiyongdengjizhengbianhao'][-5:],
    'quhuadaima':'',
    'shouliriqi':'',
    'chanpinbianhao':'',
    'shiyongdanwei':'',
    'page':'1',
    'rows':'10'
}

resp_search = session.post(url_rw,headers=headers,data=select_data)
temp = json.loads(resp_search.text)

if temp['total'] == 0:
    print('无记录')
else:
    print('查到记录')
    #print(resp_search.text)

    #获取renwuId、bjdId、baseId
    renwuId = temp['rows'][0]['renwuId']
    bjdId = temp['rows'][0]['bjdId']
    baseId = temp['rows'][0]['baseId']

    #增加报告
    annalid = add(session,headers,renwuId,bjdId,baseId)
    #增加壁厚测定报告
    add_bhcd(session,headers,annalid)

    #选择指定报告